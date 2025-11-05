# plugins/letter_plugin.py
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime

class Plugin:
    def get_tools(self):
        return [{
            "type": "function",
            "function": {
                "name": "write_letter",
                "description": "Generate a professional letter (leave, resignation, etc.)",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "type": {"type": "string", "enum": ["leave", "resignation", "recommendation"]},
                        "to": {"type": "string"},
                        "reason": {"type": "string"},
                        "dates": {"type": "string"},
                        "save": {"type": "boolean", "default": True}
                    },
                    "required": ["type", "to", "reason"]
                }
            }
        }]

    def execute(self, function_name, **args):
        if function_name != "write_letter":
            return "Invalid"

        doc = Document()
        doc.add_heading(f"{args['type'].title()} Application", 0)

        p = doc.add_paragraph()
        p.add_run("Date: ").bold = True
        p.add_run(datetime.now().strftime("%B %d, %Y"))

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("To: ").bold = True
        p.add_run(args['to'])

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Subject: ").bold = True
        p.add_run(f"{args['type'].title()} Request")

        doc.add_paragraph(args['reason'])
        if args.get('dates'):
            doc.add_paragraph(f"Dates: {args['dates']}")

        doc.add_paragraph("Thank you for your understanding.")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("[Your Name]")

        filename = f"{args['type']}_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(filename)

        return f"Letter saved as `{filename}` in project folder."